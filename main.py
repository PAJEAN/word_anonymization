from docx import Document
import logging
import os
import shutil
import time
import unicodedata

logging.basicConfig(filename='file.log', format='%(asctime)s [%(filename)s:%(lineno)d] %(message)s', level=logging.INFO)

class App:
    AnonymizedNamesFilename = 'NAMES.txt'
    ToTranslateDirPath = 'INPUT'
    TranslatedDirPath = 'OUTPUT'
    ErrorDirPath = 'ERROR'

    def __init__(self):
        pass
    
    def _AnonymizedNamesFile(self) -> list[str]:
        if not os.path.isfile(App.AnonymizedNamesFilename):
            with open(App.AnonymizedNamesFilename, 'w'):
                return []
        with open(App.AnonymizedNamesFilename, 'r') as fin:
            with open(App.AnonymizedNamesFilename, 'r', encoding='utf-8') as fin:
                anonymizedNames = [ligne.strip() for ligne in fin if ligne.strip()]
            return anonymizedNames

    def _CompactText(self, paragraphs) -> str:
        texts = ''
        for para in paragraphs:
            texts += para.text
        return texts

    def _DirGuard(self, dirpath: str):
        if not os.path.isdir(dirpath):
            os.makedirs(dirpath)
            logging.warning(f"Création du dossier {dirpath}")

    def _IsAnonymized(self, texts: str, mapping: dict[str, str]) -> tuple[bool, tuple[str, str, str]]:
        unknown = []
        is_anonymized = True
        texts_lower = texts.lower()
        textsRemoveAccent = self._RemoveAccents(texts)
        textsRemoveAccentLower = self._RemoveAccents(texts_lower)
        for name in mapping:
            if name in texts:
                is_anonymized = False
                unknown.append((name, "Le formatage du fichier word empêche l'anonymisation complète", ''))
            elif name.lower() in texts_lower:
                is_anonymized = False
                unknown.append((name, "Le nom a été trouvé dans le document avec une casse différente", texts[texts_lower.index(name.lower()): texts_lower.index(name.lower()) + len(name)]))
            elif self._RemoveAccents(name) in textsRemoveAccent:
                is_anonymized = False
                unknown.append((name, "Le nom a été trouvé dans le document avec une accentuation différente", texts[textsRemoveAccent.index(name.lower()): textsRemoveAccent.index(name.lower()) + len(name)]))
            elif self._RemoveAccents(name.lower()) in textsRemoveAccentLower:
                is_anonymized = False
                unknown.append((name, "Le nom a été trouvé dans le document avec une casse et une accentuation différentes", texts[textsRemoveAccentLower.index(name.lower()): textsRemoveAccentLower.index(name.lower()) + len(name)]))
        return is_anonymized, unknown

    def _MoveTo(self, srcFilepath: str, dstDirpath: str, filename: str):
        if os.path.exists(os.path.join(dstDirpath, filename)):
            os.remove(os.path.join(dstDirpath, filename))
        shutil.move(srcFilepath, dstDirpath)
    
    def _RemoveAccents(self, input_str: str) -> str:
        nfkd_form = unicodedata.normalize('NFKD', input_str)
        return u''.join([c for c in nfkd_form if not unicodedata.combining(c)])
    
    def _RemoveEmptyFolders(self):
        try:
            for folder in os.listdir(App.ToTranslateDirPath):
                dirpath = os.path.join(App.ToTranslateDirPath, folder)
                if os.path.isdir(dirpath):
                    folderToKeep = set()
                    walk = list(os.walk(dirpath))
                    for path, _, _ in walk[::-1]: # Down/Top.
                        files = [f_ for f_ in os.listdir(path) if os.path.isfile(f'{path}/{f_}')]
                        if len(files) == 0 and not path in folderToKeep :
                            os.rmdir(path)
                        else:
                            toKeep = []
                            current = []
                            for part in path.split('/'):
                                current.append(part)
                                toKeep.append('/'.join(current)) 
                            folderToKeep = folderToKeep.union(toKeep)
        except Exception as err:
            logging.error(err)

    def _Replace(self, paragraphs, mapping):
        for para in paragraphs:
            for name, anonymizedName in mapping.items():
                if name in para.text:
                    # Accès direct au run est nécessaire pour que la mise en forme soit conservée.
                    for run in para.runs:
                        run.text = run.text.replace(name, anonymizedName)    
                        
    def Run(self):
        self._DirGuard(App.ToTranslateDirPath)
        self._DirGuard(App.TranslatedDirPath)
        self._DirGuard(App.ErrorDirPath)

        anonymizedNames = self._AnonymizedNamesFile()
        if len(anonymizedNames) == 0:
            logging.info(f'Aucun nom à anonymiser')
            exit(f'Aucun noms à anonymiser')

        # Manage a mapping name_1 -> [ANONYME_1], etc.
        mapping = {name: f"[ANONYME_{i+1}]" for i, name in enumerate(anonymizedNames)}
        
        # Filters files by them extensions: .zip, .xlxs and .docx.
        files = [os.path.join(dp, f) for dp, dn, filenames in os.walk(App.ToTranslateDirPath) for f in filenames if os.path.splitext(f)[1] in ['.docx']]

        print(f'{"-" * 60}')
        for filepath in files:
            try:
                fileRecords =  os.path.split(filepath)
                fileRecords = fileRecords[1:] # Delete INPUT folder.
                fileRecords = list(filter(lambda curr: curr != '', fileRecords))
                filename = fileRecords[-1]

                # Check if the file is completely downloaded.
                fileSize = -1
                while fileSize != os.path.getsize(filepath):
                    fileSize = os.path.getsize(filepath)
                    time.sleep(1)

                logging.info(f'Anonymize the following file: {filepath}')
                
                doc = Document(filepath)

                self._Replace(doc.paragraphs, mapping)

                # Loop on tables (if they contain names).
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            self._Replace(cell.paragraphs, mapping)

                basename, extension = os.path.splitext(filename)
                newName = f'{basename}{extension}'
                newPath = os.path.join(App.TranslatedDirPath, newName)
                doc.save(newPath)

                doc = Document(newPath)
                
                compactText = self._CompactText(doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            compactText += self._CompactText(cell.paragraphs)

                isAnonymizedConcatXmlTag, namesConcatXmlTag = self._IsAnonymized(compactText, mapping)

                if not isAnonymizedConcatXmlTag:
                    print(f'[FAIL]  {filepath}')
                    if len(namesConcatXmlTag) > 0:
                        for n in namesConcatXmlTag:
                            print(f'            {n[0]}{f" -> {n[2]}" if n[2] != '' else ''} // {n[1]}')
                else:
                    print(f'[CHECK] {filepath}')
            except Exception as err:
                print(err)
                logging.error(err)
                # self._MoveTo(filepath, App.ErrorDirPath, filename)
        
        # Remove folders.
        self._RemoveEmptyFolders()

        print(f'{"-" * 60}')

        nb_col_1 = 0
        nb_col_2 = 0
        for key in mapping:
            if len(key) > nb_col_1:
                nb_col_1 = len(key)
            if len(mapping[key]) > nb_col_2:
                nb_col_2 = len(mapping[key])
        nb_col_1 += 2
        nb_col_2 += 2
        row_format = '{:>'+ str(nb_col_1) + '}{:>5}{:>' + str(nb_col_2) + '}'
        for key in mapping:
            print(row_format.format(key, '->', mapping[key]))
            

if __name__ == '__main__':
    app = App()
    app.Run()
    
